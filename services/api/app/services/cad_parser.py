"""
CAD Parser Service - Extract metadata from CAD/BIM files
Supports IFC, DXF, DWG (via conversion), and generates SubmissionProfile
"""

import logging
from typing import Dict, Any, Optional
from pathlib import Path
import json
import subprocess
import tempfile
import shutil
import re
import time

logger = logging.getLogger(__name__)


class IFCParser:
    """Parser for IFC (Industry Foundation Classes) files using IfcOpenShell"""
    
    def __init__(self):
        try:
            import ifcopenshell
            import ifcopenshell.util.element as element_util
            import ifcopenshell.util.shape as shape_util
            self.ifcopenshell = ifcopenshell
            self.element_util = element_util
            self.shape_util = shape_util
        except ImportError:
            logger.error("IfcOpenShell not installed. Install with: pip install ifcopenshell")
            raise
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse IFC file and extract building metadata
        
        Args:
            file_path: Path to IFC file
            
        Returns:
            Dictionary with extracted metadata
        """
        try:
            ifc_file = self.ifcopenshell.open(file_path)
            
            # Extract basic info
            project = ifc_file.by_type("IfcProject")[0] if ifc_file.by_type("IfcProject") else None
            
            metadata = {
                "file_schema": ifc_file.schema,
                "project_name": project.Name if project else "Unknown",
                "building": self._extract_building_info(ifc_file),
                "storeys": self._extract_storeys(ifc_file),
                "spaces": self._extract_spaces(ifc_file),
                "elements": self._extract_elements(ifc_file),
                "systems": self._detect_systems(ifc_file),
                "properties": self._extract_properties(ifc_file),
                "quantities": self._calculate_quantities(ifc_file),
            }
            
            logger.info(f"Successfully parsed IFC file: {file_path}")
            return metadata
            
        except Exception as e:
            logger.error(f"Error parsing IFC file: {e}")
            raise
    
    def _extract_building_info(self, ifc_file) -> Dict[str, Any]:
        """Extract building-level information"""
        buildings = ifc_file.by_type("IfcBuilding")
        
        if not buildings:
            return {"found": False}
        
        building = buildings[0]
        
        return {
            "found": True,
            "name": building.Name or "Unnamed",
            "description": building.Description or "",
            "elevation": building.ElevationOfRefHeight if hasattr(building, "ElevationOfRefHeight") else None,
            "address": self._extract_address(building),
        }
    
    def _extract_address(self, building) -> Optional[Dict[str, str]]:
        """Extract building address if available"""
        try:
            if hasattr(building, "BuildingAddress") and building.BuildingAddress:
                addr = building.BuildingAddress
                return {
                    "street": addr.AddressLines[0] if addr.AddressLines else None,
                    "city": addr.Town,
                    "region": addr.Region,
                    "postal_code": addr.PostalCode,
                    "country": addr.Country,
                }
        except Exception:
            pass
        return None
    
    def _extract_storeys(self, ifc_file) -> Dict[str, Any]:
        """Extract building storeys/floors"""
        storeys = ifc_file.by_type("IfcBuildingStorey")
        
        storey_list = []
        for storey in storeys:
            storey_list.append({
                "name": storey.Name or "Unnamed",
                "elevation": storey.Elevation if hasattr(storey, "Elevation") else None,
                "description": storey.Description or "",
            })
        
        return {
            "count": len(storeys),
            "storeys": sorted(storey_list, key=lambda x: x.get("elevation") or 0),
        }
    
    def _extract_spaces(self, ifc_file) -> Dict[str, Any]:
        """Extract spaces/rooms"""
        spaces = ifc_file.by_type("IfcSpace")
        
        space_list = []
        for space in spaces:
            space_list.append({
                "name": space.Name or "Unnamed",
                "long_name": space.LongName if hasattr(space, "LongName") else None,
                "description": space.Description or "",
            })
        
        return {
            "count": len(spaces),
            "spaces": space_list[:50],  # Limit to first 50 for summary
        }
    
    def _extract_elements(self, ifc_file) -> Dict[str, int]:
        """Count different element types"""
        element_types = [
            "IfcWall", "IfcDoor", "IfcWindow", "IfcSlab", "IfcRoof",
            "IfcStair", "IfcRailing", "IfcColumn", "IfcBeam",
            "IfcCovering", "IfcFurnishingElement"
        ]
        
        counts = {}
        for elem_type in element_types:
            elements = ifc_file.by_type(elem_type)
            if elements:
                counts[elem_type.replace("Ifc", "").lower()] = len(elements)
        
        return counts
    
    def _detect_systems(self, ifc_file) -> Dict[str, bool]:
        """Detect building systems presence"""
        systems = {
            "electrical": bool(ifc_file.by_type("IfcElectricAppliance") or 
                             ifc_file.by_type("IfcElectricDistributionBoard")),
            "plumbing": bool(ifc_file.by_type("IfcSanitaryTerminal") or
                           ifc_file.by_type("IfcPipeFitting")),
            "hvac": bool(ifc_file.by_type("IfcAirTerminal") or
                       ifc_file.by_type("IfcDuctFitting")),
            "fire_protection": bool(ifc_file.by_type("IfcFireSuppressionTerminal")),
        }
        
        return systems
    
    def _extract_properties(self, ifc_file) -> Dict[str, Any]:
        """Extract property sets"""
        property_sets = ifc_file.by_type("IfcPropertySet")
        
        properties = {}
        for pset in property_sets[:20]:  # Limit for performance
            pset_name = pset.Name
            if pset_name and pset.HasProperties:
                properties[pset_name] = len(pset.HasProperties)
        
        return {
            "property_sets_count": len(property_sets),
            "sample_property_sets": properties,
        }
    
    def _calculate_quantities(self, ifc_file) -> Dict[str, Any]:
        """Calculate basic quantities"""
        try:
            walls = ifc_file.by_type("IfcWall")
            slabs = ifc_file.by_type("IfcSlab")
            
            # Basic area estimation (simplified)
            total_wall_area = 0
            total_floor_area = 0
            
            # Note: Full quantity calculation requires shape processing
            # This is a simplified version
            
            return {
                "wall_count": len(walls),
                "slab_count": len(slabs),
                "estimated_wall_area_sqm": "requires_shape_processing",
                "estimated_floor_area_sqm": "requires_shape_processing",
            }
        except Exception as e:
            logger.warning(f"Quantity calculation failed: {e}")
            return {}


class DXFParser:
    """Parser for DXF and DWG (AutoCAD Drawing) files using ezdxf"""
    
    def __init__(self):
        try:
            import ezdxf
            self.ezdxf = ezdxf
        except ImportError:
            logger.error("ezdxf not installed. Install with: pip install ezdxf")
            raise
    
    def _clean_dxf_file(self, dxf_path: str) -> tuple[str, int]:
        """
        Clean up problematic MTEXT entities in DXF file.
        LibreDWG can put malformed formatting codes in BOTH group code positions AND text values.
        Uses a three-pass approach to identify and remove all affected entities.
        
        Args:
            dxf_path: Path to DXF file to clean
            
        Returns:
            Tuple of (cleaned_file_path, number_of_entities_removed)
        """
        try:
            cleaned_path = dxf_path.replace('.dxf', '_cleaned.dxf')
            
            logger.info(f"Starting DXF cleaning process for: {dxf_path}")
            
            with open(dxf_path, 'r', encoding='utf-8', errors='ignore') as infile:
                lines = infile.readlines()
            
            logger.info(f"Scanning {len(lines)} lines for malformed MTEXT patterns...")
            
            # PASS 1: Identify all line numbers containing malformed patterns
            malformed_lines = set()
            for line_num, line in enumerate(lines):
                has_problem = False
                
                try:
                    # Look for pipe-based formatting codes
                    if '|i0' in line or '|c0' in line or ('|p' in line and '|' in line):
                        has_problem = True
                    # Look for font codes
                    elif 'fArial' in line or 'fTimes' in line or 'fSwis' in line:
                        has_problem = True  
                    # Look for height/formatting codes with backslash
                    elif ('\\H' in line and 'x;' in line) or '\\S^' in line or '\\A1{' in line or '\\pql' in line:
                        has_problem = True
                    # Look for scientific notation in group code position (LibreDWG bug)
                    elif ('E-' in line or 'E+' in line) and (line.strip().startswith('-') or line.strip()[0].isdigit()):
                        has_problem = True
                    # Look for text names where numbers should be (GENERATED_STYLE, layer names, etc.)
                    elif line.strip() and line.strip()[0].isalpha() and '_' in line:
                        # Likely a name appearing where a group code should be
                        has_problem = True
                        
                    if has_problem:
                        malformed_lines.add(line_num)
                        if len(malformed_lines) <= 10:  # Log first 10 only
                            logger.debug(f"Found malformed pattern at line {line_num}: {line.strip()[:150]}")
                except Exception as line_error:
                    logger.warning(f"Error checking line {line_num}: {line_error}")
                    continue
            
            logger.info(f"Found {len(malformed_lines)} lines with malformed patterns")
            
            # Debug: Show some examples of malformed lines
            if malformed_lines:
                logger.info(f"Example malformed line numbers: {sorted(list(malformed_lines))[:10]}")
            
            # Note: We intentionally do NOT replace malformed lines directly as this corrupts the DXF structure
            # Instead, we rely on removing entire malformed MTEXT entities and using recovery mode
            
            # PASS 2: Map remaining malformed lines to their MTEXT entity start positions
            malformed_entities = set()
            current_entity_start = None
            in_mtext = False
            entity_depth = 0
            
            for i in range(len(lines)):
                line = lines[i].strip()
                
                # Detect MTEXT entity start
                if i + 1 < len(lines) and line == '0' and lines[i + 1].strip() == 'MTEXT':
                    current_entity_start = i
                    in_mtext = True
                    entity_depth = 0
                    logger.debug(f"MTEXT entity starts at line {i}")
                # Detect next entity start (end of current MTEXT)
                # Only end MTEXT if we see '0' followed by another entity type
                elif in_mtext and line == '0':
                    # Check if next line is a new entity type
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        # If it's another entity type (not a data value), end the MTEXT
                        if next_line in ['MTEXT', 'LINE', 'LWPOLYLINE', 'CIRCLE', 'ARC', 'TEXT', 'INSERT', 
                                         'POLYLINE', 'DIMENSION', 'HATCH', 'SPLINE', 'ELLIPSE', 'POINT',
                                         'SOLID', 'TRACE', 'ATTRIB', 'BLOCK', 'ENDBLK', 'SEQEND']:
                            if current_entity_start is not None:
                                logger.debug(f"MTEXT entity ends at line {i} (started at {current_entity_start})")
                            in_mtext = False
                            current_entity_start = None
                
                # If we're in an MTEXT entity and this line is malformed, mark the entity
                if in_mtext and current_entity_start is not None and i in malformed_lines:
                    malformed_entities.add(current_entity_start)
                    logger.debug(f"Marking MTEXT entity at line {current_entity_start} for removal (malformed line {i})")
            
            logger.info(f"Identified {len(malformed_entities)} MTEXT entities to remove")
            
            # PASS 3: Write output, skipping malformed MTEXT entities and fixing missing subclass markers
            skipped_count = 0
            fixed_count = 0
            i = 0
            
            with open(cleaned_path, 'w', encoding='utf-8') as outfile:
                while i < len(lines):
                    # Check if this is the start of a malformed MTEXT entity
                    if i in malformed_entities:
                        skipped_count += 1
                        entity_start = i
                        # Skip '0' and 'MTEXT'
                        i += 2
                        # Skip all lines of this entity until next '0' group code
                        while i < len(lines):
                            if lines[i].strip() == '0':
                                # Found start of next entity, don't skip this line
                                break
                            i += 1
                        logger.debug(f"Skipped MTEXT entity starting at line {entity_start}")
                    # Check if this is a clean MTEXT entity that might be missing subclass markers
                    elif i + 1 < len(lines) and lines[i].strip() == '0' and lines[i + 1].strip() == 'MTEXT':
                        # Write the MTEXT entity start
                        outfile.write(lines[i])     # '0'
                        outfile.write(lines[i + 1])  # 'MTEXT'
                        i += 2
                        
                        # Check if subclass markers are present in next few lines
                        has_subclass_marker = False
                        lookahead_start = i
                        lookahead_limit = min(i + 20, len(lines))  # Check next 20 lines
                        
                        for j in range(lookahead_start, lookahead_limit):
                            if lines[j].strip() == '100':
                                has_subclass_marker = True
                                break
                            if lines[j].strip() == '0':  # Hit next entity
                                break
                        
                        # If missing subclass markers, add them (required for DXF R13+)
                        if not has_subclass_marker:
                            outfile.write('100\n')
                            outfile.write('AcDbEntity\n')
                            outfile.write('100\n')
                            outfile.write('AcDbMText\n')
                            fixed_count += 1
                            logger.debug(f"Added missing subclass markers to MTEXT at line {lookahead_start - 2}")
                        
                        # Write rest of the entity
                        while i < len(lines):
                            if lines[i].strip() == '0':
                                # Next entity starts, don't write this line yet
                                break
                            outfile.write(lines[i])
                            i += 1
                    else:
                        # Write line as-is
                        outfile.write(lines[i])
                        i += 1
            
            logger.info(f"Cleaned DXF file created: {cleaned_path} (removed {skipped_count} problematic MTEXT entities, fixed {fixed_count} missing subclass markers)")
            return cleaned_path, skipped_count
            
        except Exception as e:
            logger.error(f"Error cleaning DXF file: {e}", exc_info=True)
            # Return original path with 0 count if cleaning fails
            return dxf_path, 0
    
    def _translate_dwg_with_aps(self, dwg_path: str) -> Optional[Dict[str, Any]]:
        """
        Translate DWG using Autodesk Platform Services (APS) Model Derivative API.
        Translates to SVF2 format and extracts metadata directly via API.
        
        Requires APS_CLIENT_ID and APS_CLIENT_SECRET environment variables.
        APS (formerly Forge) provides 100% reliable DWG translation supporting ALL versions.
        
        Args:
            dwg_path: Path to DWG file
            
        Returns:
            Dictionary with extracted metadata, or None if translation fails
        """
        try:
            import requests
            import base64
            import time
            from app.core.config import settings
            
            # Check for credentials
            if not settings.APS_CLIENT_ID or not settings.APS_CLIENT_SECRET:
                logger.debug("APS API credentials not configured")
                return None
            
            # Step 1: Get 2-legged OAuth2 access token
            auth_url = "https://developer.api.autodesk.com/authentication/v2/token"
            
            # Use form-encoded authentication
            auth_response = requests.post(
                auth_url,
                headers={"Content-Type": "application/x-www-form-urlencoded"},
                data={
                    "client_id": settings.APS_CLIENT_ID,
                    "client_secret": settings.APS_CLIENT_SECRET,
                    "grant_type": "client_credentials",
                    "scope": "data:read data:write data:create bucket:create bucket:read"
                },
                timeout=30
            )
            
            if auth_response.status_code != 200:
                logger.error(f"APS authentication failed: {auth_response.text}")
                return None
            
            access_token = auth_response.json()["access_token"]
            
            # Step 2: Upload DWG file using signed URL approach
            bucket_key = f"cadvisor-temp-{int(time.time())}"
            object_name = Path(dwg_path).name
            
            # Create bucket using OSS v2 API (still supported for bucket creation)
            bucket_url = "https://developer.api.autodesk.com/oss/v2/buckets"
            bucket_response = requests.post(
                bucket_url,
                headers={
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json"
                },
                json={"bucketKey": bucket_key, "policyKey": "temporary"},
                timeout=30
            )
            
            # Bucket might already exist, that's OK
            if bucket_response.status_code not in [200, 201, 409]:
                logger.error(f"APS bucket creation failed: {bucket_response.text}")
                return None
                
            # Get signed upload URL
            signed_url_endpoint = f"https://developer.api.autodesk.com/oss/v2/buckets/{bucket_key}/objects/{object_name}/signeds3upload"
            signed_response = requests.get(
                signed_url_endpoint,
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=30
            )
            
            if signed_response.status_code != 200:
                logger.error(f"Failed to get signed URL: {signed_response.text}")
                return None
                
            upload_data = signed_response.json()
            upload_url = upload_data["urls"][0]
            upload_key = upload_data["uploadKey"]
            
            # Upload file to signed S3 URL
            with open(dwg_path, "rb") as f:
                s3_response = requests.put(
                    upload_url,
                    data=f,
                    timeout=300
                )
                
                if s3_response.status_code not in [200, 201]:
                    logger.error(f"S3 upload failed: {s3_response.status_code}")
                    return None
            
            # Complete the upload
            complete_url = f"https://developer.api.autodesk.com/oss/v2/buckets/{bucket_key}/objects/{object_name}/signeds3upload"
            complete_response = requests.post(
                complete_url,
                headers={
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json"
                },
                json={"uploadKey": upload_key},
                timeout=30
            )
            
            if complete_response.status_code not in [200, 201]:
                logger.error(f"Failed to complete upload: {complete_response.text}")
                return None
            
            # Generate URN from objectId
            object_id = complete_response.json()["objectId"]
            urn = base64.urlsafe_b64encode(object_id.encode()).decode().rstrip("=")
            
            # Step 3: Trigger Model Derivative translation job (DWG → SVF2)
            derivative_url = "https://developer.api.autodesk.com/modelderivative/v2/designdata/job"
            job_response = requests.post(
                derivative_url,
                headers={
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json",
                    "x-ads-force": "true"  # Force new translation
                },
                json={
                    "input": {"urn": urn},
                    "output": {
                        "formats": [{"type": "svf2", "views": ["2d", "3d"]}]
                    }
                },
                timeout=30
            )
            
            if job_response.status_code not in [200, 201]:
                logger.error(f"APS translation job failed: {job_response.text}")
                return None
            
            # Step 4: Poll for completion (max 5 minutes)
            manifest_url = f"https://developer.api.autodesk.com/modelderivative/v2/designdata/{urn}/manifest"
            max_wait = 300  # 5 minutes
            start_time = time.time()
            
            while time.time() - start_time < max_wait:
                manifest_response = requests.get(
                    manifest_url,
                    headers={"Authorization": f"Bearer {access_token}"},
                    timeout=30
                )
                
                if manifest_response.status_code == 200:
                    manifest = manifest_response.json()
                    status = manifest.get("status")
                    progress = manifest.get("progress", "")
                    
                    logger.debug(f"APS translation status: {status}, progress: {progress}")
                    
                    if status == "success":
                        logger.info(f"APS translation succeeded, extracting metadata via API")
                        
                        # Step 5: Extract metadata via APS Metadata API (Enhanced with fallback)
                        metadata = self._extract_metadata_from_aps_enhanced(urn, access_token)
                        
                        if metadata:
                            logger.info(f"Successfully extracted metadata from APS (enhanced)")
                            return metadata
                        
                        # Fallback to legacy method if enhanced fails
                        logger.info("Enhanced extraction failed, trying legacy method...")
                        metadata = self._extract_metadata_from_aps_legacy(urn, access_token)
                        
                        if metadata:
                            logger.info(f"Successfully extracted metadata from APS (legacy)")
                            return metadata
                        else:
                            logger.error("Failed to extract metadata from APS (both methods)")
                            return None
                    
                    elif status == "failed":
                        logger.error(f"APS conversion failed: {progress}")
                        return None
                    
                    elif status == "inprogress":
                        # Continue polling
                        pass
                
                time.sleep(5)
            
            logger.error("APS translation timed out after 5 minutes")
            return None
            
        except ImportError:
            logger.debug("requests library not available for APS API")
            return None
        except Exception as e:
            logger.error(f"APS API error: {e}")
            return None
    
    def _extract_metadata_from_aps_enhanced(self, urn: str, access_token: str) -> Optional[Dict[str, Any]]:
        """
        Extract metadata from APS using enhanced Object Tree and Query APIs.
        This provides hierarchical structure and more granular data extraction.
        
        Uses:
        1. Object Tree API - Gets hierarchical BIM structure with parent-child relationships
        2. Properties API - Gets detailed properties for all objects
        
        Args:
            urn: Base64-encoded URN of the translated file
            access_token: APS OAuth2 access token
            
        Returns:
            Dictionary with extracted metadata or None
        """
        try:
            import requests
            
            # Step 1: List model views (get viewable GUIDs)
            metadata_url = f"https://developer.api.autodesk.com/modelderivative/v2/designdata/{urn}/metadata"
            views_response = requests.get(
                metadata_url,
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=30
            )
            
            if views_response.status_code != 200:
                logger.error(f"Failed to get model views: {views_response.text}")
                return None
            
            views_data = views_response.json()
            metadata_list = views_data.get("data", {}).get("metadata", [])
            
            if not metadata_list:
                logger.warning("No model views found in APS metadata")
                return {"views": [], "objects": []}
            
            # Step 2: Extract Object Tree (hierarchical structure) and Properties from all views
            all_objects = []
            all_hierarchy = []
            all_raw_responses = []  # Store complete unfiltered API responses
            view_info = []
            
            for view in metadata_list:
                guid = view.get("guid")
                view_name = view.get("name", "Unnamed")
                view_role = view.get("role", "unknown")
                
                if not guid:
                    continue
                
                view_info.append({
                    "guid": guid,
                    "name": view_name,
                    "role": view_role
                })
                
                # Step 2a: Get Object Tree (hierarchical structure) for this view
                tree_url = f"https://developer.api.autodesk.com/modelderivative/v2/designdata/{urn}/metadata/{guid}?forceget=true"
                tree_response = requests.get(
                    tree_url,
                    headers={
                        "Authorization": f"Bearer {access_token}",
                        "x-ads-force": "true"  # Force retrieval even on recoverable failures
                    },
                    timeout=120  # Increased timeout for large files
                )
                
                if tree_response.status_code == 200:
                    tree_data = tree_response.json()
                    hierarchy = tree_data.get("data", {}).get("objects", [])
                    node_count = self._count_tree_objects(hierarchy)
                    logger.info(f"✓ Extracted object tree from view '{view_name}' with {node_count} objects")
                    all_hierarchy.extend(hierarchy)
                    # Store raw tree response
                    all_raw_responses.append({
                        "view_guid": guid,
                        "view_name": view_name,
                        "api_type": "object_tree",
                        "response": tree_data
                    })
                elif tree_response.status_code == 202:
                    logger.info(f"⟲ Object tree still processing for view '{view_name}', will use properties API")
                elif tree_response.status_code == 413:
                    logger.warning(f"⚠ Object tree too large for view '{view_name}' (>800MB), using properties only")
                else:
                    logger.warning(f"✗ Failed to get object tree for view '{view_name}': {tree_response.status_code} - {tree_response.text[:200]}")
                
                # Step 2b: Get all properties for this view with retry logic
                properties_url = f"https://developer.api.autodesk.com/modelderivative/v2/designdata/{urn}/metadata/{guid}/properties?forceget=true"
                
                # Retry up to 6 times (30 seconds total) for properties to be ready
                max_retries = 6
                retry_delay = 5
                
                for attempt in range(max_retries):
                    props_response = requests.get(
                        properties_url,
                        headers={
                            "Authorization": f"Bearer {access_token}",
                            "x-ads-force": "true",  # Force retrieval even on recoverable failures
                            "Accept-Encoding": "gzip"  # Request compression for large responses
                        },
                        timeout=120  # Increased timeout for large files
                    )
                    
                    if props_response.status_code == 200:
                        props_data = props_response.json()
                        objects = props_data.get("data", {}).get("collection", [])
                        
                        # Store complete raw properties response (unfiltered)
                        all_raw_responses.append({
                            "view_guid": guid,
                            "view_name": view_name,
                            "api_type": "properties",
                            "response": props_data,
                            "object_count": len(objects)
                        })
                        
                        # Add view context to each object
                        for obj in objects:
                            obj["view_name"] = view_name
                            obj["view_role"] = view_role
                        
                        all_objects.extend(objects)
                        logger.info(f"✓ Extracted {len(objects)} objects from view '{view_name}' (response size: {len(props_response.content)} bytes)")
                        break  # Success, exit retry loop
                        
                    elif props_response.status_code == 202:
                        # Properties still processing
                        if attempt < max_retries - 1:
                            logger.info(f"⟲ Properties still processing for view '{view_name}', retrying in {retry_delay}s (attempt {attempt + 1}/{max_retries})...")
                            time.sleep(retry_delay)
                        else:
                            logger.warning(f"⚠ Properties for view '{view_name}' still not ready after {max_retries} attempts")
                    elif props_response.status_code == 413:
                        logger.error(f"✗ Properties too large for view '{view_name}' (>800MB) - consider using Query API for filtering")
                        break
                    else:
                        logger.warning(f"✗ Failed to get properties for view '{view_name}': {props_response.status_code} - {props_response.text[:200]}")
                        break  # Non-retryable error
            
            logger.info(f"📊 Total extraction: {len(all_objects)} objects from {len(view_info)} views, hierarchy nodes: {self._count_tree_objects(all_hierarchy)}")
            logger.info(f"💾 Stored {len(all_raw_responses)} raw API responses (complete unfiltered data)")
            
            # Step 3: Organize and structure the metadata with hierarchy
            return self._structure_aps_metadata_enhanced(all_objects, view_info, all_hierarchy, all_raw_responses)
            
        except Exception as e:
            logger.error(f"Error extracting APS metadata (enhanced): {e}")
            return None
    
    def _count_tree_objects(self, tree_nodes: list) -> int:
        """
        Recursively count all objects in the hierarchy tree.
        
        Args:
            tree_nodes: List of tree node objects
            
        Returns:
            Total count of objects in tree
        """
        count = len(tree_nodes)
        for node in tree_nodes:
            if "objects" in node:
                count += self._count_tree_objects(node["objects"])
        return count
    
    def _structure_aps_metadata_enhanced(self, objects: list, views: list, hierarchy: list, raw_responses: list = None) -> Dict[str, Any]:
        """
        Structure APS metadata with hierarchical information (ENHANCED VERSION).
        
        Args:
            objects: List of objects with properties from APS
            views: List of model views
            hierarchy: Hierarchical object tree structure
            raw_responses: Complete unfiltered API responses from APS
            
        Returns:
            Structured metadata dictionary with hierarchy information
        """
        # First get base structure from existing method
        structured_metadata = self._structure_aps_metadata(objects, views)
        
        # Enhance with hierarchy information
        if hierarchy:
            structured_metadata["hierarchy"] = {
                "available": True,
                "tree": hierarchy[:100],  # Limit to first 100 nodes for database storage
                "total_nodes": self._count_tree_objects(hierarchy),
                "note": "Object hierarchy showing parent-child relationships and BIM structure"
            }
            structured_metadata["extraction_method"] = "aps_enhanced_api (tree + properties)"
        
        # Store complete raw APS responses (unfiltered, all data preserved)
        if raw_responses:
            structured_metadata["aps_raw_responses"] = {
                "available": True,
                "response_count": len(raw_responses),
                "responses": raw_responses,
                "note": "Complete unfiltered API responses from APS - ALL data preserved without truncation",
                "warning": "This contains the complete dataset which may be very large"
            }
        
        return structured_metadata
    
    def _extract_metadata_from_aps_legacy(self, urn: str, access_token: str) -> Optional[Dict[str, Any]]:
        """
        Extract metadata from APS using legacy flat properties approach (BACKUP METHOD).
        Uses the Metadata API to get properties of all objects in flat structure.
        
        This is kept as a fallback if the enhanced Object Tree API fails.
        
        Args:
            urn: Base64-encoded URN of the translated file
            access_token: APS OAuth2 access token
            
        Returns:
            Dictionary with extracted metadata or None
        """
        try:
            import requests
            
            # Step 1: List model views (get viewable GUIDs)
            metadata_url = f"https://developer.api.autodesk.com/modelderivative/v2/designdata/{urn}/metadata"
            views_response = requests.get(
                metadata_url,
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=30
            )
            
            if views_response.status_code != 200:
                logger.error(f"Failed to get model views (legacy): {views_response.text}")
                return None
            
            views_data = views_response.json()
            metadata_list = views_data.get("data", {}).get("metadata", [])
            
            if not metadata_list:
                logger.warning("No model views found in APS metadata (legacy)")
                return {"views": [], "objects": []}
            
            # Step 2: Extract properties from all views (flat structure only)
            all_objects = []
            all_raw_responses = []  # Store complete unfiltered API responses
            view_info = []
            
            for view in metadata_list:
                guid = view.get("guid")
                view_name = view.get("name", "Unnamed")
                view_role = view.get("role", "unknown")
                
                if not guid:
                    continue
                
                view_info.append({
                    "guid": guid,
                    "name": view_name,
                    "role": view_role
                })
                
                # Get all properties for this view with retry logic (LEGACY METHOD WITH FORCEGET)
                properties_url = f"https://developer.api.autodesk.com/modelderivative/v2/designdata/{urn}/metadata/{guid}/properties?forceget=true"
                
                # Retry up to 6 times (30 seconds total) for properties to be ready
                max_retries = 6
                retry_delay = 5
                
                for attempt in range(max_retries):
                    props_response = requests.get(
                        properties_url,
                        headers={
                            "Authorization": f"Bearer {access_token}",
                            "x-ads-force": "true",  # Force retrieval even on recoverable failures
                            "Accept-Encoding": "gzip"  # Request compression
                        },
                        timeout=120  # Increased timeout
                    )
                    
                    if props_response.status_code == 200:
                        props_data = props_response.json()
                        objects = props_data.get("data", {}).get("collection", [])
                        
                        # Store complete raw properties response (unfiltered)
                        all_raw_responses.append({
                            "view_guid": guid,
                            "view_name": view_name,
                            "api_type": "properties",
                            "response": props_data,
                            "object_count": len(objects)
                        })
                        
                        # Add view context to each object
                        for obj in objects:
                            obj["view_name"] = view_name
                            obj["view_role"] = view_role
                        
                        all_objects.extend(objects)
                        logger.info(f"✓ [LEGACY] Extracted {len(objects)} objects from view '{view_name}' (response: {len(props_response.content)} bytes)")
                        break  # Success, exit retry loop
                        
                    elif props_response.status_code == 202:
                        # Properties still processing
                        if attempt < max_retries - 1:
                            logger.info(f"⟲ [LEGACY] Properties still processing for view '{view_name}', retrying in {retry_delay}s (attempt {attempt + 1}/{max_retries})...")
                            time.sleep(retry_delay)
                        else:
                            logger.warning(f"⚠ [LEGACY] Properties for view '{view_name}' still not ready after {max_retries} attempts")
                    elif props_response.status_code == 413:
                        logger.error(f"✗ [LEGACY] Properties too large for view '{view_name}' (>800MB)")
                        break
                    else:
                        logger.warning(f"✗ [LEGACY] Failed to get properties for view '{view_name}': {props_response.status_code} - {props_response.text[:200]}")
                        break  # Non-retryable error
            
            logger.info(f"📊 [LEGACY] Total extraction: {len(all_objects)} objects from {len(view_info)} views")
            logger.info(f"💾 [LEGACY] Stored {len(all_raw_responses)} raw API responses")
            
            # Step 3: Organize and structure the metadata (without hierarchy)
            structured = self._structure_aps_metadata(all_objects, view_info)
            
            # Add raw responses to legacy method as well
            if all_raw_responses:
                structured["aps_raw_responses"] = {
                    "available": True,
                    "response_count": len(all_raw_responses),
                    "responses": all_raw_responses,
                    "note": "Complete unfiltered API responses from APS (legacy method)"
                }
            
            return structured
            
        except Exception as e:
            logger.error(f"Error extracting APS metadata (legacy): {e}")
            return None
    
    def _structure_aps_metadata(self, objects: list, views: list) -> Dict[str, Any]:
        """
        Structure APS metadata into a format compatible with our database schema and frontend.
        
        Args:
            objects: List of objects with properties from APS
            views: List of model views
            
        Returns:
            Structured metadata dictionary matching frontend expectations
        """
        # Extract entities by type and layer information
        entity_types = {}
        layer_dict = {}
        all_properties = []
        
        for obj in objects:
            obj_name = obj.get("name", "Unknown")
            obj_id = obj.get("objectid")
            props = obj.get("properties", {})
            
            # Map APS object names to DXF-like entity types
            # Extract entity type from name (e.g., "Line [ABC]" -> "LINE")
            entity_type = obj_name.split('[')[0].strip().upper()
            if not entity_type:
                entity_type = "UNKNOWN"
            
            entity_types[entity_type] = entity_types.get(entity_type, 0) + 1
            
            # Extract layer information from properties
            general_props = props.get("General", {})
            layer_name = general_props.get("Layer")
            
            if layer_name:
                if layer_name not in layer_dict:
                    layer_dict[layer_name] = {
                        "name": layer_name,
                        "color": general_props.get("Color"),
                        "linetype": general_props.get("Linetype"),
                        "lineweight": general_props.get("Lineweight")
                    }
            
            # Collect all properties for raw data tab
            all_properties.append({
                "objectid": obj_id,
                "name": obj_name,
                "type": entity_type,
                "properties": props,
                "view": obj.get("view"),
                "role": obj.get("role")
            })
        
        # Convert layers dict to list
        layers_list = list(layer_dict.values())
        
        # Build structured response matching frontend expectations
        structured_metadata = {
            # Status fields (used by Parsing Report tab)
            "processing_status": "completed",
            "source_format": "dwg",
            "extraction_method": "aps_metadata_api",
            
            # Entities field (used by Geometry tab - expects "total" field)
            "entities": {
                "total": len(objects),
                **entity_types  # Spread individual entity type counts
            },
            
            # Layers field (used by Building tab - expects "count" and "layers" array)
            "layers": {
                "count": len(layers_list),
                "layers": layers_list[:100]  # Limit to first 100 for display
            },
            
            # Views from APS (2D/3D model views)
            "views": {
                "count": len(views),
                "views": views
            },
            
            # Raw objects data (for advanced analysis and Raw Data tab)
            "objects": {
                "total_count": len(objects),
                "objects": all_properties[:500]  # Limit to first 500 for database storage
            },
            
            # Additional metadata
            "metadata_note": "Extracted via Autodesk APS Model Derivative Metadata API with full CAD properties",
            "aps_extraction": True
        }

        # Extract compliance-relevant data from ALL objects (runs on full untruncated list)
        compliance = self._extract_compliance_analysis(objects)
        structured_metadata["rooms"] = {
            "count": len(compliance["rooms"]),
            "rooms": compliance["rooms"]
        }
        structured_metadata["fire_elements"] = {
            "count": len(compliance["fire_elements"]),
            "items": compliance["fire_elements"]
        }
        structured_metadata["evacuation"] = compliance["evacuation"]
        structured_metadata["text_annotations"] = {
            "count": len(compliance["text_annotations"]),
            "items": compliance["text_annotations"]
        }
        structured_metadata["structural_elements"] = compliance["structural_elements"]

        return structured_metadata

    def _decode_mtext(self, text: str) -> str:
        """Strip AutoCAD MTEXT formatting codes to produce plain text."""
        if not text:
            return ""
        # Remove formatting escapes: \A1; \pql; \fArial|b0|i0|c0|p34; \W0.9x; \pqc; \pqr; etc.
        t = re.sub(r'\\[A-Za-z][^;\\{}\n]*;', '', text)
        # \P = paragraph break
        t = re.sub(r'\\P', ' ', t)
        # Strip braces
        t = re.sub(r'[{}]', '', t)
        # Tab → space
        t = t.replace('\t', ' ')
        # Collapse whitespace
        t = re.sub(r' {2,}', ' ', t).strip()
        return t

    def _extract_compliance_analysis(self, objects: list) -> Dict[str, Any]:
        """
        Scan all APS objects (untruncated) for compliance-relevant data.
        Extracts rooms, fire ratings, evacuation distances, text annotations,
        and structural element counts.

        Returns dict with keys: rooms, fire_elements, evacuation,
                                text_annotations, structural_elements
        """
        rei_pat  = re.compile(r'REI[\s\-]?(\d+)', re.IGNORECASE)
        ei_pat   = re.compile(r'EI[\s\-]?(\d+)(?:[\-\u2013]([A-Z]+))?', re.IGNORECASE)
        evac_pat = re.compile(r'lungime\s+de\s+evacuare\s*=\s*([\d.,]+)\s*(m|mm)?', re.IGNORECASE)
        fire_kw  = [
            'rezistent la foc', 'rezistenta la foc', 'incombustibil',
            'desfumare', 'evacuare', ' rei ', ' ei ', 'rezistenta foc'
        ]

        rooms            = []
        fire_elements    = []
        evacuation_info  = []
        text_annotations = []
        beam_count       = 0
        slab_count       = 0
        structural_layers: set = set()

        seen_fire_texts   = set()
        seen_room_handles = set()

        for obj in objects:
            obj_name   = (obj.get('name') or '').strip()
            props      = obj.get('properties') or {}
            if not isinstance(props, dict):
                continue

            general   = props.get('General') or {}
            layer_val = (general.get('Layer') or '').strip()
            handle    = (general.get('Handle') or '').strip()
            layer_l   = layer_val.lower()
            name_l    = obj_name.lower()

            # ── Rooms from Attributes property set ─────────────────────────
            attrs     = props.get('Attributes') or {}
            room_name = (attrs.get('ROOM_NAME') or '').strip()
            if room_name and handle not in seen_room_handles:
                seen_room_handles.add(handle)
                height_raw = (attrs.get('sCustom1Value') or '').strip()
                height_m: Optional[float] = None
                hm = re.search(r'[Hh]\s*=\s*([\d.,]+)', height_raw)
                if hm:
                    try:
                        height_m = float(hm.group(1).replace(',', '.'))
                    except ValueError:
                        pass
                rooms.append({
                    "name": room_name,
                    "floor_code": (attrs.get('floorCode') or '').strip(),
                    "floor_covering": (attrs.get('stFloorCovering') or '').strip(),
                    "height_m": height_m,
                    "area_marker": (attrs.get('measuredAreaCode') or '').strip(),
                    "layer": layer_val,
                })

            # ── Structural element counting ────────────────────────────────
            if 'beam' in name_l or 'grinda' in layer_l:
                beam_count += 1
            if 'placa' in name_l or 'slab' in name_l or 'plac' in layer_l:
                slab_count += 1
            if any(x in layer_l for x in ('str', 'structura', 'grinda', 'ifc model', 'plac')):
                structural_layers.add(layer_val)

            # ── Text / MTEXT content ───────────────────────────────────────
            text_pset   = props.get('Text') or {}
            raw_content = (text_pset.get('Contents') or '').strip()
            if not raw_content:
                continue

            decoded   = self._decode_mtext(raw_content)
            decoded_l = decoded.lower()
            if len(decoded) < 3:
                continue

            text_annotations.append({
                "text": decoded[:500],
                "layer": layer_val,
                "text_height_mm": (text_pset.get('Text height') or '').strip(),
                "object_name": obj_name,
            })

            # ── Fire elements ──────────────────────────────────────────────
            has_rei  = rei_pat.search(decoded)
            has_ei   = ei_pat.search(decoded)
            is_fire  = has_rei or has_ei or any(kw in decoded_l for kw in fire_kw)
            has_evac = evac_pat.search(decoded)

            if is_fire and not has_evac:
                key = decoded[:120]
                if key not in seen_fire_texts:
                    seen_fire_texts.add(key)

                    etype = 'general'
                    if any(x in decoded_l for x in ('planseu', 'planşeu', 'plansee', 'placa')):
                        etype = 'floor_slab'
                    elif any(x in decoded_l for x in ('perete', 'zid')):
                        etype = 'wall'
                    elif 'scara' in decoded_l or 'scări' in decoded_l:
                        etype = 'stair'
                    elif any(x in decoded_l for x in ('usa ', 'uşa ', 'ușa ')):
                        etype = 'door'
                    elif 'incombustibil' in decoded_l:
                        etype = 'insulation'
                    elif 'desfumare' in decoded_l:
                        etype = 'smoke_ventilation'

                    ratings = (
                        [f"REI {m.group(1)}" for m in rei_pat.finditer(decoded)] +
                        [f"EI {m.group(1)}{('-' + m.group(2)) if m.group(2) else ''}"
                         for m in ei_pat.finditer(decoded)]
                    )
                    fire_elements.append({
                        "element_type": etype,
                        "rating": ratings[0] if len(ratings) == 1 else (', '.join(ratings) if ratings else None),
                        "all_ratings": ratings,
                        "text": decoded[:300],
                        "layer": layer_val,
                        "source": "mtext",
                    })

            # ── Evacuation distances ───────────────────────────────────────
            if has_evac:
                for m in evac_pat.finditer(decoded):
                    try:
                        val = float(m.group(1).replace(',', '.'))
                    except ValueError:
                        continue
                    unit = (m.group(2) or 'm').lower()
                    evacuation_info.append({
                        "type": "travel_distance",
                        "value": val,
                        "unit": unit,
                        "text": decoded[:200],
                        "layer": layer_val,
                    })

        return {
            "rooms": rooms,
            "fire_elements": fire_elements,
            "evacuation": evacuation_info,
            "text_annotations": text_annotations[:300],   # cap at 300 items for DB size
            "structural_elements": {
                "beams": beam_count,
                "slabs": slab_count,
                "structural_layers": sorted(structural_layers),
            },
        }

    def _convert_dwg_with_libredwg(self, dwg_path: str) -> Optional[str]:
        """
        Convert DWG to DXF using LibreDWG's dwg2dxf utility (fallback).
        Less reliable but works for simple drawings.
        
        Args:
            dwg_path: Path to DWG file
            
        Returns:
            Path to converted DXF file, or None if conversion fails
        """
        try:
            temp_dir = tempfile.gettempdir()
            dxf_path = Path(temp_dir) / f"{Path(dwg_path).stem}_libredwg.dxf"
            
            # Run dwg2dxf conversion
            result = subprocess.run(
                ["dwg2dxf", "-o", str(dxf_path), dwg_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0 and dxf_path.exists():
                logger.info(f"LibreDWG conversion succeeded: {dxf_path}")
                
                # Clean up problematic MTEXT entities
                logger.info("Cleaning DXF file to remove problematic MTEXT entities...")
                cleaned_path, cleaned_count = self._clean_dxf_file(str(dxf_path))
                
                if cleaned_count > 0:
                    logger.info(f"Removed {cleaned_count} problematic MTEXT entities")
                
                return cleaned_path
            else:
                logger.warning(f"LibreDWG conversion failed: {result.stderr}")
                return None
                
        except subprocess.TimeoutExpired:
            logger.error("LibreDWG conversion timed out after 60 seconds")
            return None
        except FileNotFoundError:
            logger.debug("dwg2dxf (LibreDWG) not found")
            return None
        except Exception as e:
            logger.warning(f"LibreDWG conversion error: {e}")
            return None

    def _process_dwg_file(self, dwg_path: str) -> Optional[Dict[str, Any]]:
        """
        Process DWG file using best available method.
        Tries multiple methods in order of reliability:
        1. Autodesk Platform Services (APS) - most reliable, 100% compatibility, direct metadata extraction
        2. LibreDWG conversion to DXF - free fallback, ~60% success rate
        
        Args:
            dwg_path: Path to DWG file
            
        Returns:
            Dictionary with extracted metadata, or None if all methods fail
        """
        logger.info(f"Processing DWG file: {dwg_path}")
        
        # Method 1: Try Autodesk Platform Services (APS) first (best quality, direct metadata)
        logger.info("Attempting translation with Autodesk Platform Services (APS)...")
        result = self._translate_dwg_with_aps(dwg_path)
        if result:
            return result
        
        # Method 2: Fallback to LibreDWG + ezdxf parsing (free but less reliable)
        logger.info("APS not available, attempting conversion with LibreDWG...")
        dxf_path = self._convert_dwg_with_libredwg(dwg_path)
        if dxf_path:
            # Parse the DXF file with ezdxf
            logger.info(f"Parsing converted DXF file: {dxf_path}")
            return self._parse_dxf_file(dxf_path, source_format="dwg")
        
        # All methods failed
        logger.error("All DWG processing methods failed")
        return None
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse DXF or DWG file and extract drawing metadata
        
        Args:
            file_path: Path to DXF or DWG file
            
        Returns:
            Dictionary with extracted metadata
        """
        try:
            # Check if file is DWG or DXF
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.dwg':
                logger.info(f"DWG file detected, processing: {file_path}")
                # Process DWG directly (APS translates to SVF2 + metadata, or LibreDWG converts to DXF)
                result = self._process_dwg_file(file_path)
                
                if not result:
                    return {
                        "error": "DWG processing failed",
                        "processing_status": "failed",
                        "message": "Could not process DWG file. All methods failed.",
                        "source_format": "dwg"
                    }
                
                return result
            
            # For DXF files, parse directly with ezdxf
            return self._parse_dxf_file(file_path, source_format="dxf")
            
        except Exception as e:
            logger.error(f"Error parsing CAD file: {e}")
            return {
                "error": str(e),
                "processing_status": "failed"
            }
    
    def _parse_dxf_file(self, file_path: str, source_format: str = "dxf") -> Dict[str, Any]:
        """
        Parse DXF file with ezdxf and extract metadata.
        
        Args:
            file_path: Path to DXF file
            source_format: Original format ("dxf" or "dwg")
            
        Returns:
            Dictionary with extracted metadata
        """
        parsing_log = []
        converted_path = None
        
        try:
            parsing_log.append(f"Parsing DXF file (source: {source_format})")
            
            # Parse DXF file with ezdxf (use recovery mode for robustness)
            parsing_log.append("Attempting standard ezdxf parsing...")
            try:
                doc = self.ezdxf.readfile(file_path)
                parsing_log.append("Standard parsing successful")
            except Exception as read_error:
                error_msg = str(read_error)
                parsing_log.append(f"Standard parsing failed: {error_msg[:200]}")
                logger.warning(f"Standard read failed, trying recovery mode: {read_error}")
                
                # Try recovery mode for malformed DXF files
                parsing_log.append("Attempting ezdxf recovery mode...")
                try:
                    from ezdxf import recover
                    doc, auditor = recover.readfile(file_path)
                    if auditor.has_errors:
                        parsing_log.append(f"Recovery mode found {len(auditor.errors)} errors but recovered")
                        logger.warning(f"DXF file has {len(auditor.errors)} errors (recovered)")
                    parsing_log.append("Recovery mode successful")
                    logger.info("Successfully recovered DXF file using recovery mode")
                except Exception as recovery_error:
                    parsing_log.append(f"Recovery mode failed: {str(recovery_error)[:200]}")
                    logger.error(f"Recovery mode also failed: {recovery_error}")
                    
                    # Last resort: try to extract text directly from DXF file
                    logger.info("Attempting raw text extraction from DXF file...")
                    parsing_log.append("Attempting raw text extraction as last resort...")
                    try:
                        raw_text = self._extract_raw_text_from_dxf(file_path)
                        parsing_log.append(f"Raw text extraction found {len(raw_text)} characters")
                        return {
                            "processing_status": "partial",
                            "error": str(read_error),
                            "recovery_error": str(recovery_error),
                            "source_format": source_format,
                            "message": "File has parsing errors, but text content was extracted",
                            "raw_text_content": raw_text[:10000],  # Limit to first 10KB
                            "text_extraction_method": "raw_dxf_parsing",
                            "parsing_log": parsing_log
                        }
                    except Exception as text_error:
                        parsing_log.append(f"Raw text extraction failed: {str(text_error)}")
                        logger.error(f"Raw text extraction also failed: {text_error}")
                        # If everything fails, return error
                        return {
                            "processing_status": "failed",
                            "error": str(read_error),
                            "recovery_error": str(recovery_error),
                            "source_format": source_format,
                            "message": "File parsing failed. The file may be corrupted, use an unsupported DWG version, or contain invalid data.",
                            "parsing_log": parsing_log
                        }
            
            parsing_log.append("Extracting metadata from parsed document...")
            
            # Count total entities
            entities = self._count_entities(doc)
            total_entities = sum(entities.values()) if entities else 0
            
            metadata = {
                "processing_status": "completed",
                "source_format": source_format,
                "extraction_method": "ezdxf_local_parsing",
                "dxf_version": doc.dxfversion,
                "layers": self._extract_layers(doc),
                "blocks": self._extract_blocks(doc),
                "entities": entities,
                "text_annotations": self._extract_text(doc),
                "dimensions": self._extract_dimensions(doc),
                "viewport_info": self._extract_viewport_info(doc),
                "parsing_log": parsing_log
            }
            
            parsing_log.append(f"Metadata extraction complete: {len(metadata.get('layers', {}).get('layers', []))} layers, {total_entities} entities")
            logger.info(f"Successfully parsed DXF file: {file_path}")
            return metadata
            
        except Exception as e:
            logger.error(f"Error parsing CAD file: {e}")
            return {
                "error": str(e),
                "processing_status": "failed"
            }
        finally:
            # Clean up temporary converted file
            if converted_path and Path(converted_path).exists():
                try:
                    Path(converted_path).unlink()
                    logger.debug(f"Cleaned up temporary DXF file: {converted_path}")
                except Exception as e:
                    logger.warning(f"Failed to clean up temporary file: {e}")
    
    def _extract_raw_text_from_dxf(self, file_path: str) -> str:
        """
        Extract text content directly from DXF file by parsing the text file
        This is a fallback when ezdxf parsing fails
        """
        try:
            texts = []
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()
                
                # Simple state machine to extract TEXT and MTEXT content
                in_text = False
                in_mtext = False
                text_buffer = []
                
                for i, line in enumerate(lines):
                    line = line.strip()
                    
                    # Detect TEXT or MTEXT entities
                    if line in ['TEXT', 'MTEXT']:
                        in_text = True
                        text_buffer = []
                        continue
                    
                    # Group code 1 is the text content
                    if in_text and i > 0 and lines[i-1].strip() == '1':
                        text_content = line
                        if text_content and len(text_content) > 2:
                            texts.append(text_content)
                        in_text = False
                        continue
                    
                    # Group code 3 is MTEXT content (can be multi-line)
                    if in_text and i > 0 and lines[i-1].strip() == '3':
                        text_content = line
                        if text_content and len(text_content) > 2:
                            text_buffer.append(text_content)
                        continue
                    
                    # Reset state on new entity
                    if line in ['ENDSEC', 'ENDBLK'] or (line.isdigit() and line == '0'):
                        if text_buffer:
                            texts.append(' '.join(text_buffer))
                        in_text = False
                        text_buffer = []
            
            logger.info(f"Extracted {len(texts)} text entities using raw parsing")
            return '\n'.join(texts)
            
        except Exception as e:
            logger.error(f"Raw text extraction failed: {e}")
            return ""
    
    def _extract_layers(self, doc) -> Dict[str, Any]:
        """Extract layer information"""
        layers = []
        
        for layer in doc.layers:
            layers.append({
                "name": layer.dxf.name,
                "color": layer.dxf.color,
                "linetype": layer.dxf.linetype,
                "is_locked": layer.is_locked(),
                "is_off": layer.is_off(),
            })
        
        return {
            "count": len(layers),
            "layers": layers,
        }
    
    def _extract_blocks(self, doc) -> Dict[str, Any]:
        """Extract block definitions"""
        blocks = []
        
        for block in doc.blocks:
            if not block.name.startswith("*"):  # Skip anonymous blocks
                blocks.append({
                    "name": block.name,
                    "entity_count": len(block),
                })
        
        return {
            "count": len(blocks),
            "blocks": blocks[:50],  # Limit for summary
        }
    
    def _count_entities(self, doc) -> Dict[str, int]:
        """Count different entity types"""
        msp = doc.modelspace()
        
        entity_counts = {}
        for entity in msp:
            entity_type = entity.dxftype()
            entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1
        
        return entity_counts
    
    def _extract_text(self, doc) -> Dict[str, Any]:
        """Extract text annotations"""
        msp = doc.modelspace()
        
        texts = []
        for entity in msp.query("TEXT MTEXT"):
            text_content = entity.dxf.text if hasattr(entity.dxf, "text") else str(entity)
            if text_content:
                texts.append({
                    "content": text_content[:100],  # Truncate long text
                    "layer": entity.dxf.layer,
                })
        
        return {
            "count": len(texts),
            "sample_texts": texts[:20],  # First 20 for analysis
        }
    
    def _extract_dimensions(self, doc) -> Dict[str, Any]:
        """Extract dimension entities"""
        msp = doc.modelspace()
        
        dimensions = list(msp.query("DIMENSION"))
        
        return {
            "count": len(dimensions),
            "has_dimensions": len(dimensions) > 0,
        }
    
    def _extract_viewport_info(self, doc) -> Dict[str, Any]:
        """Extract viewport and layout information"""
        layouts = []
        
        for layout in doc.layouts:
            layouts.append({
                "name": layout.name,
                "entity_count": len(layout),
            })
        
        return {
            "layout_count": len(layouts),
            "layouts": layouts,
        }


class DocumentParser:
    """Parser for PDF, DOCX, and other document formats"""
    
    def __init__(self):
        self.pdf_parser = PDFParser()
        self.docx_parser = DOCXParser()
    
    def parse(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Parse document based on MIME type
        
        Args:
            file_path: Path to document
            mime_type: MIME type of document
            
        Returns:
            Dictionary with extracted content
        """
        if mime_type == "application/pdf":
            return self.pdf_parser.parse(file_path)
        elif "wordprocessing" in mime_type or mime_type == "application/msword":
            return self.docx_parser.parse(file_path)
        elif mime_type == "text/plain":
            return self._parse_text(file_path)
        else:
            return {"error": f"Unsupported document type: {mime_type}"}
    
    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                "type": "text",
                "content": content,
                "length": len(content),
                "lines": len(content.split('\n')),
            }
        except Exception as e:
            logger.error(f"Error parsing text file: {e}")
            return {"error": str(e)}


class PDFParser:
    """
    Advanced PDF parser with OCR, table detection, and spatial analysis.
    Optimized for fire safety compliance document analysis.
    """
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Extract comprehensive content from PDF including text, tables, images, legends.
        Performs OCR on images and scanned pages with Romanian support.
        """
        try:
            import fitz  # PyMuPDF
            import pytesseract
            from PIL import Image
            import io
            import numpy as np
            import cv2
            
            logger.info(f"Starting advanced PDF parsing: {file_path}")
            parsing_log = [
                "Detected file type: PDF",
                "Starting advanced PDF content extraction with OCR..."
            ]
            
            # Open PDF
            doc = fitz.open(file_path)
            total_pages = len(doc)
            parsing_log.append(f"PDF contains {total_pages} pages")
            
            # Initialize extraction results
            all_text = []
            all_blocks = []  # Text blocks with position
            tables = []
            images_extracted = []
            legends_found = []
            color_codes = []
            page_metadata = []
            
            # Romanian fire safety keywords
            fire_safety_keywords = [
                'rezistenta la foc', 'rezistență la foc', 'REI', 'EI', 'compartimentare',
                'evacuare', 'iesire', 'ieșire', 'scara', 'legenda', 'legendă',
                'ignifug', 'protectie', 'protecție', 'rezistent', 'incendiu'
            ]
            
            # Process each page
            for page_num in range(total_pages):
                page = doc[page_num]
                page_info = {
                    'page_number': page_num + 1,
                    'width': page.rect.width,
                    'height': page.rect.height
                }
                
                # Extract text with position information
                blocks = page.get_text("dict")["blocks"]
                page_has_text = False
                
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                text_content = span.get("text", "").strip()
                                if text_content:
                                    page_has_text = True
                                    all_text.append(text_content)
                                    all_blocks.append({
                                        'text': text_content,
                                        'page': page_num + 1,
                                        'bbox': span.get("bbox"),  # (x0, y0, x1, y1)
                                        'font': span.get("font"),
                                        'size': span.get("size"),
                                        'color': span.get("color")
                                    })
                
                # If page has no/little text, it might be scanned - OCR the whole page
                if not page_has_text or len([b for b in all_blocks if b['page'] == page_num + 1]) < 5:
                    try:
                        # Render page to image
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        
                        # Perform OCR
                        ocr_text = pytesseract.image_to_string(img, lang='ron+eng', config='--psm 3')
                        
                        if ocr_text.strip():
                            all_text.append(ocr_text)
                            parsing_log.append(f"Page {page_num + 1}: Performed full-page OCR (scanned/low text)")
                    except Exception as ocr_error:
                        logger.debug(f"Full-page OCR failed for page {page_num + 1}: {ocr_error}")
                
                # Extract tables using text block positions
                page_tables = self._extract_tables_from_blocks(blocks, page_num + 1)
                tables.extend(page_tables)
                
                # Extract images and perform OCR
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # Convert to PIL Image
                        pil_image = Image.open(io.BytesIO(image_bytes))
                        
                        # Perform OCR with Romanian support
                        ocr_text = pytesseract.image_to_string(
                            pil_image,
                            lang='ron+eng',
                            config='--psm 6'
                        )
                        
                        if ocr_text.strip():
                            all_text.append(ocr_text)
                            images_extracted.append({
                                'page': page_num + 1,
                                'index': img_index,
                                'format': image_ext,
                                'ocr_text': ocr_text,
                                'size': len(image_bytes)
                            })
                            
                            # Check for fire safety legends in images
                            if any(keyword in ocr_text.lower() for keyword in fire_safety_keywords):
                                legends_found.append({
                                    'page': page_num + 1,
                                    'source': 'image_ocr',
                                    'content': ocr_text[:500]  # First 500 chars
                                })
                        
                        # Detect color codes using OpenCV
                        img_array = np.array(pil_image.convert('RGB'))
                        img_bgr = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                        
                        # Find dominant colors
                        dominant_colors = self._extract_dominant_colors(img_bgr)
                        if dominant_colors:
                            color_codes.extend([{
                                'page': page_num + 1,
                                'rgb': color,
                                'percentage': percentage
                            } for color, percentage in dominant_colors])
                    
                    except Exception as img_error:
                        logger.debug(f"Error processing image {img_index} on page {page_num + 1}: {img_error}")
                        continue
                
                # Detect legends and tables in text
                page_text = " ".join([b['text'] for b in all_blocks if b['page'] == page_num + 1])
                
                # Look for legend patterns
                legend_patterns = [
                    r'legend[aă][\s:]+.*rezisten[țt][aă].*foc',
                    r'tabel.*rezisten[țt][aă].*foc',
                    r'simbol.*REI.*\d+',
                    r'nota[țt]ii.*REI'
                ]
                
                for pattern in legend_patterns:
                    matches = list(re.finditer(pattern, page_text, re.IGNORECASE | re.DOTALL))[:5]  # Max 5 per pattern
                    for match in matches:
                        # Extract context around match
                        start = max(0, match.start() - 100)
                        end = min(len(page_text), match.end() + 200)
                        context = page_text[start:end]
                        
                        legends_found.append({
                            'page': page_num + 1,
                            'source': 'text_pattern',
                            'content': context
                        })
                
                page_metadata.append(page_info)
            
            doc.close()
            
            # Combine all text
            full_text = "\n".join(all_text)
            
            # Extract fire safety specific information
            rei_codes = self._extract_rei_codes(full_text)
            compartments = self._detect_compartmentation(full_text)
            egress_routes = self._detect_egress_routes(full_text, all_blocks)
            
            parsing_log.append(f"Extracted {len(full_text)} characters of text")
            parsing_log.append(f"Found {len(all_blocks)} text blocks with position data")
            parsing_log.append(f"Extracted {len(images_extracted)} images with OCR")
            parsing_log.append(f"Detected {len(tables)} tables")
            parsing_log.append(f"Found {len(legends_found)} potential fire safety legends")
            parsing_log.append(f"Identified {len(rei_codes)} REI classifications")
            parsing_log.append(f"Detected {len(compartments)} compartmentation references")
            parsing_log.append(f"Found {len(egress_routes)} egress route indicators")
            parsing_log.append("PDF parsing completed successfully")
            
            return {
                'type': 'pdf',
                'success': True,
                'processing_status': 'completed',
                'pages': total_pages,
                'text': full_text,  # For backward compatibility
                'text_content': full_text,
                'char_count': len(full_text),
                'text_blocks': all_blocks[:1000],  # Limit for performance
                'tables': tables,
                'images': images_extracted,
                'legends': legends_found,
                'color_codes': color_codes[:50],  # Top 50 colors
                'fire_safety_data': {
                    'rei_codes': rei_codes,
                    'compartments': compartments,
                    'egress_routes': egress_routes
                },
                'page_metadata': page_metadata,
                'parsing_log': parsing_log,
                'text_extraction_method': 'pymupdf_ocr'
            }
            
        except ImportError as e:
            logger.error(f"Required library not installed: {e}")
            return {
                'type': 'pdf',
                'success': False,
                'error': f"Required library not installed: {e}",
                'message': 'PDF parsing requires PyMuPDF, pytesseract, PIL, opencv-python',
                'processing_status': 'failed',
                'parsing_log': ['PDF parsing failed: Missing required libraries'],
                'text': ''
            }
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}", exc_info=True)
            return {
                'type': 'pdf',
                'success': False,
                'error': str(e),
                'message': f'Failed to parse PDF: {str(e)}',
                'processing_status': 'failed',
                'parsing_log': [f'PDF parsing failed: {str(e)}'],
                'text': ''
            }
    
    def _extract_tables_from_blocks(self, blocks: list, page_num: int) -> list:
        """Extract table-like structures from text blocks based on spatial positioning"""
        tables = []
        
        # Group blocks by approximate Y coordinate (rows)
        rows = {}
        for block in blocks:
            if block.get("type") != 0:
                continue
            
            for line in block.get("lines", []):
                bbox = line.get("bbox")
                if not bbox:
                    continue
                
                y_coord = int(bbox[1] / 10) * 10  # Round to nearest 10 for grouping
                if y_coord not in rows:
                    rows[y_coord] = []
                
                text = " ".join([span.get("text", "") for span in line.get("spans", [])])
                rows[y_coord].append({
                    'text': text.strip(),
                    'x': bbox[0],
                    'y': bbox[1]
                })
        
        # Detect tables (rows with similar number of columns)
        row_list = sorted(rows.items())
        if len(row_list) > 3:  # At least 3 rows for a table
            # Look for consecutive rows with multiple columns
            current_table = []
            for y_coord, cells in row_list:
                if len(cells) >= 2:  # At least 2 columns
                    sorted_cells = sorted(cells, key=lambda c: c['x'])
                    current_table.append([cell['text'] for cell in sorted_cells])
                else:
                    if len(current_table) >= 3:
                        tables.append({
                            'page': page_num,
                            'rows': len(current_table),
                            'data': current_table
                        })
                    current_table = []
            
            if len(current_table) >= 3:
                tables.append({
                    'page': page_num,
                    'rows': len(current_table),
                    'data': current_table
                })
        
        return tables
    
    def _extract_dominant_colors(self, img_bgr, num_colors: int = 5) -> list:
        """Extract dominant colors from image using K-means clustering"""
        try:
            import cv2
            import numpy as np
            
            # Resize for faster processing
            img_small = cv2.resize(img_bgr, (100, 100))
            pixels = img_small.reshape(-1, 3).astype(float)
            
            # Use K-means to find dominant colors
            criteria = (cv2.TERM_CRITERIA_EPS + cv2.TERM_CRITERIA_MAX_ITER, 10, 1.0)
            _, labels, centers = cv2.kmeans(pixels, num_colors, None, criteria, 10, cv2.KMEANS_PP_CENTERS)
            
            # Calculate color percentages
            unique, counts = np.unique(labels, return_counts=True)
            color_percentages = counts / counts.sum()
            
            # Return colors sorted by frequency
            colors_with_percentage = [
                (tuple(int(c) for c in centers[i]), float(color_percentages[i]))
                for i in range(len(centers))
            ]
            return sorted(colors_with_percentage, key=lambda x: x[1], reverse=True)
        
        except Exception as e:
            logger.debug(f"Error extracting colors: {e}")
            return []
    
    def _extract_rei_codes(self, text: str) -> list:
        """Extract REI fire resistance codes from text"""
        rei_pattern = r'REI[-:\s=]*(\d+)'
        matches = re.finditer(rei_pattern, text, re.IGNORECASE)
        
        rei_codes = []
        seen_codes = set()
        for match in matches:
            code = match.group(0)
            if code not in seen_codes:
                seen_codes.add(code)
                rei_codes.append({
                    'code': code,
                    'minutes': int(match.group(1)),
                    'context': text[max(0, match.start()-50):min(len(text), match.end()+50)]
                })
        
        return rei_codes[:100]  # Limit to first 100
    
    def _detect_compartmentation(self, text: str) -> list:
        """Detect fire compartmentation references"""
        compartment_patterns = [
            r'compartiment[^\n]{0,100}foc',
            r'sector[^\n]{0,100}incendiu',
            r'separare[^\n]{0,100}ignifug'
        ]
        
        compartments = []
        seen = set()
        for pattern in compartment_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                text_found = match.group(0)
                if text_found not in seen:
                    seen.add(text_found)
                    compartments.append(text_found)
        
        return compartments[:50]  # Limit to first 50
    
    def _detect_egress_routes(self, text: str, blocks: list) -> list:
        """Detect emergency egress route indicators"""
        egress_keywords = [
            'evacuare', 'ieșire', 'iesire', 'urgență', 'urgenta',
            'scară', 'scara', 'cale', 'rută', 'ruta'
        ]
        
        egress_routes = []
        for block in blocks:
            block_text = block.get('text', '').lower()
            if any(keyword in block_text for keyword in egress_keywords):
                egress_routes.append({
                    'text': block['text'],
                    'page': block.get('page'),
                    'position': block.get('bbox')
                })
        
        return egress_routes[:50]  # Limit to first 50


class DOCXParser:
    """DOCX document extraction"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract full text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            combined_text = "\n\n".join(paragraphs)
            
            return {
                "type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "text": combined_text,
                "char_count": len(combined_text)
            }
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return {"error": str(e), "type": "docx", "text": ""}


class CADParserService:
    """Main service for parsing CAD and document files"""
    
    def __init__(self):
        self.ifc_parser = None
        self.dxf_parser = None
        self.doc_parser = DocumentParser()
    
    def parse_file(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Parse file based on type
        
        Args:
            file_path: Path to file
            mime_type: MIME type
            
        Returns:
            Parsed metadata
        """
        try:
            # IFC files
            if "ifc" in mime_type.lower() or file_path.endswith('.ifc'):
                if not self.ifc_parser:
                    self.ifc_parser = IFCParser()
                return {"type": "ifc", "data": self.ifc_parser.parse(file_path)}
            
            # DXF and DWG files
            elif "dxf" in mime_type.lower() or file_path.endswith('.dxf'):
                if not self.dxf_parser:
                    self.dxf_parser = DXFParser()
                return {"type": "dxf", "data": self.dxf_parser.parse(file_path)}
            
            elif "dwg" in mime_type.lower() or "acad" in mime_type.lower() or file_path.endswith('.dwg'):
                if not self.dxf_parser:
                    self.dxf_parser = DXFParser()
                return {"type": "dwg", "data": self.dxf_parser.parse(file_path)}
            
            # Documents
            elif mime_type in ["application/pdf", "application/msword", "text/plain"] or \
                 "wordprocessing" in mime_type:
                return {"type": "document", "data": self.doc_parser.parse(file_path, mime_type)}
            
            else:
                return {"type": "unsupported", "mime_type": mime_type}
                
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            return {"type": "error", "error": str(e)}
